#!/usr/bin/env python3
import os
from pathlib import Path
from docx import Document

def create_test_aoa():
    doc = Document()
    
    title = doc.add_heading('ARTICLES OF ASSOCIATION', 0)
    title.alignment = 1
    
    doc.add_heading('1. COMPANY INFORMATION', level=1)
    doc.add_paragraph('Company Name: TechStart Innovations LLC')
    doc.add_paragraph('Jurisdiction: UAE Federal Courts')  # RED FLAG
    
    doc.add_heading('2. SHARE CAPITAL', level=1)
    doc.add_paragraph('The company may have share capital of AED 100,000.')  # RED FLAG
    
    doc.add_heading('3. DIRECTORS', level=1)
    doc.add_paragraph('The company possibly shall have directors.')  # RED FLAG
    
    doc.add_heading('4. GOVERNING LAW', level=1)
    doc.add_paragraph('These Articles shall be governed by Dubai Courts.')  # RED FLAG
    
    return doc

def create_compliant_aoa():
    doc = Document()
    
    title = doc.add_heading('ARTICLES OF ASSOCIATION', 0)
    title.alignment = 1
    
    doc.add_heading('1. COMPANY INFORMATION', level=1)
    doc.add_paragraph('Company Name: ADGM Tech Solutions Ltd')
    doc.add_paragraph('Jurisdiction: Abu Dhabi Global Market (ADGM)')  # CORRECT
    
    doc.add_heading('2. SHARE CAPITAL', level=1)
    doc.add_paragraph('The authorized share capital shall be AED 100,000.')  # CORRECT
    
    doc.add_heading('3. DIRECTORS', level=1)
    doc.add_paragraph('The Company shall have at least one director.')  # CORRECT
    
    doc.add_heading('4. GOVERNING LAW', level=1)
    doc.add_paragraph('These Articles shall be governed by ADGM laws and subject to ADGM Courts.')  # CORRECT
    
    doc.add_heading('5. SIGNATURES', level=1)
    doc.add_paragraph('IN WITNESS WHEREOF, the parties have executed these Articles.')
    doc.add_paragraph('_' * 30)
    doc.add_paragraph('Signature: _______________')
    doc.add_paragraph('Date: _______________')
    
    return doc

def main():
    test_dir = Path('data/test_documents')
    test_dir.mkdir(parents=True, exist_ok=True)
    
    print('📄 Creating test documents...')
    
    docs = [
        ('articles_with_issues.docx', create_test_aoa),
        ('compliant_articles.docx', create_compliant_aoa)
    ]
    
    for filename, generator in docs:
        doc = generator()
        filepath = test_dir / filename
        doc.save(filepath)
        print(f'✅ Created: {filepath}')
    
    print('\n🎯 Test documents ready!')
    print('💡 Upload them in the Streamlit app to test compliance checking')

if __name__ == '__main__':
    main()
