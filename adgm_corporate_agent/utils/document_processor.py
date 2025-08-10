"""
Document Processing Utilities
Handles DOCX parsing, modification, and inline commenting.
"""

import io
import asyncio
from typing import Dict, List, Any, Optional, Tuple
from pathlib import Path
import hashlib
import re

from docx import Document
from docx.shared import RGBColor, Inches
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.shared import OxmlElement, qn
import mammoth

from adgm_corporate_agent.utils.logger import setup_logger
from adgm_corporate_agent.utils.cache_manager import get_cache_manager

logger = setup_logger(__name__)
cache = get_cache_manager()

class AdvancedDocumentProcessor:
    """
    Advanced document processor with multi-modal parsing and intelligent commenting.
    Supports DOCX manipulation, structure analysis, and inline commenting.
    """
    
    def __init__(self):
        """Initialize the document processor."""
        self.supported_formats = ['.docx']
        self.max_file_size = 10 * 1024 * 1024  # 10MB
        
        # Document type patterns for classification
        self.doc_type_patterns = {
            "Articles of Association": [
                r"articles\s+of\s+association",
                r"company\s+constitution",
                r"article\s+\d+",
                r"share\s+capital",
                r"directors?\s+powers"
            ],
            "Memorandum of Association": [
                r"memorandum\s+of\s+association",
                r"objects?\s+of\s+the\s+company",
                r"liability\s+clause",
                r"capital\s+clause"
            ],
            "Board Resolution": [
                r"board\s+resolution",
                r"resolved\s+that",
                r"directors?\s+present",
                r"quorum\s+present",
                r"meeting\s+of\s+the\s+board"
            ],
            "Shareholder Resolution": [
                r"shareholder\s+resolution",
                r"general\s+meeting",
                r"shareholders?\s+present",
                r"ordinary\s+resolution",
                r"special\s+resolution"
            ],
            "Incorporation Application": [
                r"incorporation\s+application",
                r"company\s+name",
                r"registered\s+office",
                r"nature\s+of\s+business",
                r"application\s+form"
            ],
            "UBO Declaration": [
                r"ultimate\s+beneficial\s+owner",
                r"ubo\s+declaration",
                r"beneficial\s+ownership",
                r"controlling\s+interest"
            ],
            "Register of Members": [
                r"register\s+of\s+members",
                r"register\s+of\s+directors",
                r"member\s+details",
                r"share\s+register",
                r"director\s+information"
            ]
        }
        
        logger.info("Document processor initialized")

    async def parse_document(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """
        Parse document and extract structured information.
        
        Args:
            file_content: Raw file content
            filename: Original filename
            
        Returns:
            Dictionary containing parsed document data
        """
        
        # Generate document fingerprint for caching
        fingerprint = hashlib.sha256(file_content).hexdigest()
        cache_key = f"parsed_doc_{fingerprint}"
        
        # Check cache first
        cached_result = cache.get(cache_key)
        if cached_result:
            logger.info(f"Using cached parse result for {filename}")
            return cached_result
        
        try:
            # Parse with python-docx
            doc = Document(io.BytesIO(file_content))
            
            # Extract text content
            full_text = self._extract_full_text(doc)
            
            # Extract structure
            structure = self._analyze_document_structure(doc)
            
            # Classify document type
            document_type = self._classify_document_type(full_text)
            
            # Extract metadata
            metadata = self._extract_metadata(doc, filename)
            
            # Extract tables if any
            tables = self._extract_tables(doc)
            
            # Build result
            result = {
                "filename": filename,
                "document_type": document_type,
                "full_text": full_text,
                "structure": structure,
                "metadata": metadata,
                "tables": tables,
                "word_count": len(full_text.split()),
                "paragraph_count": len(doc.paragraphs),
                "fingerprint": fingerprint
            }
            
            # Cache the result
            cache.set(cache_key, result, ttl=3600)  # Cache for 1 hour
            
            logger.info(f"Successfully parsed document: {filename} (Type: {document_type})")
            return result
            
        except Exception as e:
            logger.error(f"Failed to parse document {filename}: {str(e)}")
            raise

    def _extract_full_text(self, doc: Document) -> str:
        """Extract all text content from document."""
        text_parts = []
        
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text.strip())
        
        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text.strip())
        
        return "\n".join(text_parts)

    def _analyze_document_structure(self, doc: Document) -> Dict[str, Any]:
        """Analyze document structure including headings, sections, etc."""
        structure = {
            "headings": [],
            "sections": [],
            "numbered_items": [],
            "signatures": []
        }
        
        current_section = None
        
        for i, paragraph in enumerate(doc.paragraphs):
            text = paragraph.text.strip()
            
            if not text:
                continue
            
            # Check for headings (based on style or formatting)
            if self._is_heading(paragraph):
                heading_info = {
                    "text": text,
                    "level": self._get_heading_level(paragraph),
                    "paragraph_index": i
                }
                structure["headings"].append(heading_info)
                
                # Start new section
                current_section = {
                    "title": text,
                    "start_paragraph": i,
                    "content": []
                }
                structure["sections"].append(current_section)
            
            # Check for numbered items (articles, clauses, etc.)
            if self._is_numbered_item(text):
                structure["numbered_items"].append({
                    "text": text,
                    "number": self._extract_number(text),
                    "paragraph_index": i
                })
            
            # Check for signature blocks
            if self._is_signature_block(text):
                structure["signatures"].append({
                    "text": text,
                    "paragraph_index": i
                })
            
            # Add to current section
            if current_section:
                current_section["content"].append({
                    "paragraph_index": i,
                    "text": text
                })
        
        return structure

    def _is_heading(self, paragraph) -> bool:
        """Determine if paragraph is a heading."""
        # Check style name
        style_name = paragraph.style.name.lower()
        if 'heading' in style_name:
            return True
        
        # Check formatting (bold, larger font, etc.)
        if paragraph.runs:
            first_run = paragraph.runs[0]
            if first_run.bold and len(paragraph.text) < 100:
                return True
        
        # Check for common heading patterns
        text = paragraph.text.strip()
        if re.match(r'^[A-Z\s]+, text) and len(text) < 80:
            return True
        
        return False

    def _get_heading_level(self, paragraph) -> int:
        """Get heading level (1-6)."""
        style_name = paragraph.style.name.lower()
        
        # Extract level from style name
        level_match = re.search(r'heading\s*(\d+)', style_name)
        if level_match:
            return int(level_match.group(1))
        
        # Default to level 1
        return 1

    def _is_numbered_item(self, text: str) -> bool:
        """Check if text represents a numbered item."""
        patterns = [
            r'^\d+\.',           # 1.
            r'^\d+\.\d+',        # 1.1
            r'^\(\d+\)',         # (1)
            r'^[a-z]\)',         # a)
            r'^[ivx]+\.',        # i., ii., iii.
            r'^Article\s+\d+',   # Article 1
            r'^Section\s+\d+',   # Section 1
            r'^Clause\s+\d+'     # Clause 1
        ]
        
        return any(re.match(pattern, text, re.IGNORECASE) for pattern in patterns)

    def _extract_number(self, text: str) -> str:
        """Extract numbering from text."""
        # Find first number or letter sequence
        match = re.search(r'(\d+(?:\.\d+)*|[a-z]|[ivx]+)', text, re.IGNORECASE)
        return match.group(1) if match else ""

    def _is_signature_block(self, text: str) -> bool:
        """Check if text represents a signature block."""
        signature_patterns = [
            r'signature',
            r'signed\s+by',
            r'director',
            r'secretary',
            r'witness',
            r'date.*signed',
            r'_+.*_+',  # Signature lines
            r'name:',
            r'title:',
            r'date:'
        ]
        
        return any(re.search(pattern, text, re.IGNORECASE) for pattern in signature_patterns)

    def _classify_document_type(self, text: str) -> str:
        """Classify document type based on content patterns."""
        text_lower = text.lower()
        
        # Score each document type
        scores = {}
        
        for doc_type, patterns in self.doc_type_patterns.items():
            score = 0
            for pattern in patterns:
                matches = len(re.findall(pattern, text_lower, re.IGNORECASE))
                score += matches
            scores[doc_type] = score
        
        # Return type with highest score
        if scores:
            best_type = max(scores, key=scores.get)
            if scores[best_type] > 0:
                return best_type
        
        return "Unknown Document Type"

    def _extract_metadata(self, doc: Document, filename: str) -> Dict[str, Any]:
        """Extract document metadata."""
        metadata = {
            "filename": filename,
            "file_size": 0,  # Will be set by caller
            "creation_date": None,
            "last_modified": None,
            "author": None,
            "title": None,
            "subject": None,
            "keywords": None
        }
        
        # Extract core properties if available
        try:
            core_props = doc.core_properties
            metadata.update({
                "author": core_props.author,
                "title": core_props.title,
                "subject": core_props.subject,
                "keywords": core_props.keywords,
                "creation_date": core_props.created.isoformat() if core_props.created else None,
                "last_modified": core_props.modified.isoformat() if core_props.modified else None
            })
        except:
            pass
        
        return metadata

    def _extract_tables(self, doc: Document) -> List[Dict[str, Any]]:
        """Extract table data from document."""
        tables_data = []
        
        for i, table in enumerate(doc.tables):
            table_data = {
                "table_index": i,
                "rows": len(table.rows),
                "columns": len(table.columns) if table.rows else 0,
                "data": []
            }
            
            for row in table.rows:
                row_data = []
                for cell in row.cells:
                    row_data.append(cell.text.strip())
                table_data["data"].append(row_data)
            
            tables_data.append(table_data)
        
        return tables_data

    async def add_inline_comments(self, 
                                 file_content: bytes, 
                                 comments: List[Dict[str, Any]],
                                 filename: str) -> bytes:
        """
        Add inline comments to DOCX document.
        
        Args:
            file_content: Original document content
            comments: List of comments to add
            filename: Original filename
            
        Returns:
            Modified document as bytes
        """
        try:
            # Load document
            doc = Document(io.BytesIO(file_content))
            
            # Sort comments by position (descending to avoid index shifts)
            sorted_comments = sorted(comments, key=lambda x: x.get('paragraph_index', 0), reverse=True)
            
            for comment in sorted_comments:
                await self._add_single_comment(doc, comment)
            
            # Save modified document
            output = io.BytesIO()
            doc.save(output)
            output.seek(0)
            
            logger.info(f"Added {len(comments)} comments to {filename}")
            return output.getvalue()
            
        except Exception as e:
            logger.error(f"Failed to add comments to {filename}: {str(e)}")
            raise

    async def _add_single_comment(self, doc: Document, comment: Dict[str, Any]):
        """Add a single comment to the document."""
        try:
            paragraph_index = comment.get('paragraph_index', 0)
            comment_text = comment.get('text', '')
            severity = comment.get('severity', 'Medium')
            suggestion = comment.get('suggestion', '')
            
            # Ensure paragraph index is valid
            if paragraph_index >= len(doc.paragraphs):
                paragraph_index = len(doc.paragraphs) - 1
            
            target_paragraph = doc.paragraphs[paragraph_index]
            
            # Create comment paragraph
            comment_paragraph = target_paragraph._element.getparent().insert(
                paragraph_index + 1, target_paragraph._element
            )
            
            # Format comment based on severity
            color = self._get_severity_color(severity)
            
            # Add comment text
            comment_run = comment_paragraph.add_run(f"\n[COMMENT - {severity}]: {comment_text}")
            comment_run.font.color.rgb = color
            comment_run.font.bold = True
            comment_run.font.italic = True
            
            # Add suggestion if provided
            if suggestion:
                suggestion_run = comment_paragraph.add_run(f"\nSUGGESTION: {suggestion}")
                suggestion_run.font.color.rgb = RGBColor(0, 100, 0)  # Dark green
                suggestion_run.font.italic = True
            
            # Highlight the original text if specified
            if comment.get('highlight_text'):
                self._highlight_text(target_paragraph, comment['highlight_text'], severity)
                
        except Exception as e:
            logger.error(f"Failed to add single comment: {str(e)}")

    def _get_severity_color(self, severity: str) -> RGBColor:
        """Get color based on severity level."""
        colors = {
            'Low': RGBColor(255, 165, 0),     # Orange
            'Medium': RGBColor(255, 140, 0),   # Dark orange
            'High': RGBColor(255, 69, 0),      # Red orange
            'Critical': RGBColor(220, 20, 60)  # Crimson
        }
        return colors.get(severity, RGBColor(128, 128, 128))  # Gray default

    def _highlight_text(self, paragraph, text_to_highlight: str, severity: str):
        """Highlight specific text in paragraph."""
        try:
            # Find and highlight text
            for run in paragraph.runs:
                if text_to_highlight.lower() in run.text.lower():
                    # Set highlight color based on severity
                    if severity == 'Critical':
                        run.font.highlight_color = WD_COLOR_INDEX.RED
                    elif severity == 'High':
                        run.font.highlight_color = WD_COLOR_INDEX.YELLOW
                    elif severity == 'Medium':
                        run.font.highlight_color = WD_COLOR_INDEX.TURQUOISE
                    else:
                        run.font.highlight_color = WD_COLOR_INDEX.GRAY_25
                        
        except Exception as e:
            logger.error(f"Failed to highlight text: {str(e)}")

    def validate_file(self, file_content: bytes, filename: str) -> Dict[str, Any]:
        """Validate uploaded file."""
        validation_result = {
            "valid": True,
            "errors": [],
            "warnings": [],
            "file_info": {}
        }
        
        # Check file size
        file_size = len(file_content)
        if file_size > self.max_file_size:
            validation_result["valid"] = False
            validation_result["errors"].append(
                f"File size ({file_size/1024/1024:.1f}MB) exceeds maximum allowed size ({self.max_file_size/1024/1024}MB)"
            )
        
        # Check file extension
        file_ext = Path(filename).suffix.lower()
        if file_ext not in self.supported_formats:
            validation_result["valid"] = False
            validation_result["errors"].append(
                f"Unsupported file format: {file_ext}. Supported formats: {', '.join(self.supported_formats)}"
            )
        
        # Try to open as DOCX
        try:
            doc = Document(io.BytesIO(file_content))
            validation_result["file_info"] = {
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables),
                "size_mb": round(file_size / 1024 / 1024, 2)
            }
        except Exception as e:
            validation_result["valid"] = False
            validation_result["errors"].append(f"Cannot read document: {str(e)}")
        
        return validation_result

    async def extract_key_sections(self, parsed_doc: Dict[str, Any]) -> Dict[str, str]:
        """Extract key sections from parsed document for validation."""
        sections = {}
        
        document_type = parsed_doc.get("document_type", "")
        full_text = parsed_doc.get("full_text", "")
        structure = parsed_doc.get("structure", {})
        
        # Extract sections based on document type
        if "Articles of Association" in document_type:
            sections.update(self._extract_aoa_sections(full_text, structure))
        elif "Memorandum of Association" in document_type:
            sections.update(self._extract_moa_sections(full_text, structure))
        elif "Resolution" in document_type:
            sections.update(self._extract_resolution_sections(full_text, structure))
        
        return sections

    def _extract_aoa_sections(self, text: str, structure: Dict) -> Dict[str, str]:
        """Extract key sections from Articles of Association."""
        sections = {}
        
        # Common AoA sections to extract
        section_patterns = {
            "jurisdiction": r"jurisdiction[^.]*courts?[^.]*",
            "registered_office": r"registered\s+office[^.]*",
            "share_capital": r"share\s+capital[^.]*",
            "directors_powers": r"directors?[^.]*powers?[^.]*",
            "meetings": r"meetings?[^.]*shareholders?[^.]*"
        }
        
        for section_name, pattern in section_patterns.items():
            match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
            if match:
                sections[section_name] = match.group(0)
        
        return sections

    def _extract_moa_sections(self, text: str, structure: Dict) -> Dict[str, str]:
        """Extract key sections from Memorandum of Association."""
        sections = {}
        
        section_patterns = {
            "company_name": r"name\s+of\s+the\s+company[^.]*",
            "objects": r"objects?\s+of\s+the\s+company[^.]*",
            "liability": r"liability[^.]*members?[^.]*",
            "capital": r"capital[^.]*divided[^.]*"
        }
        
        for section_name, pattern in section_patterns.items():
            match = re.search(pattern, text, re.IGNORECASE | re.DOTALL)
            if match:
                sections[section_name] = match.group(0)
        
        return sections

    def _extract_resolution_sections(self, text: str, structure: Dict) -> Dict[str, str]:
        """Extract key sections from Resolution documents."""
        sections = {}
        
        section_patterns = {
            "meeting_details": r"meeting\s+of[^.]*",
            "quorum": r"quorum[^.]*present[^.]*",
            "resolutions": r"resolved\s+that[^.]*",
            "signatures": r"signed[^.]*director[^.]*"
        }
        
        for section_name, pattern in section_patterns.items():
            matches = re.findall(pattern, text, re.IGNORECASE | re.DOTALL)
            if matches:
                sections[section_name] = " | ".join(matches)
        
        return sections

    def get_processing_stats(self) -> Dict[str, Any]:
        """Get document processing statistics."""
        # This would typically track processing metrics
        return {
            "total_processed": 0,  # Would be tracked in a real implementation
            "avg_processing_time": 0,
            "supported_formats": self.supported_formats,
            "max_file_size_mb": self.max_file_size / 1024 / 1024
        }

# Global processor instance
_processor = None

def get_document_processor() -> AdvancedDocumentProcessor:
    """Get global document processor instance."""
    global _processor
    if _processor is None:
        _processor = AdvancedDocumentProcessor()
    return _processor