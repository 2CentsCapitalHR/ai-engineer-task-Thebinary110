import os
import re
from typing import Dict, List, Any, Optional
from pathlib import Path
from datetime import datetime

from docx import Document
from docx.shared import RGBColor, Pt
import docx2txt

from ..utils.logger import setup_logger

logger = setup_logger(__name__)

class DocxProcessor:
    def __init__(self):
        self.supported_formats = ['.docx']
        
    def parse_document(self, file_path: str) -> Dict[str, Any]:
        logger.info(f'í³„ Parsing document: {file_path}')
        
        if not self._validate_file(file_path):
            raise ValueError(f'Invalid file: {file_path}')
        
        try:
            doc = Document(file_path)
            text_content = docx2txt.process(file_path)
            
            result = {
                'file_path': file_path,
                'file_name': Path(file_path).name,
                'text': text_content,
                'word_count': len(text_content.split()),
                'paragraph_count': len(doc.paragraphs),
                'parsing_timestamp': datetime.now().isoformat()
            }
            
            logger.info(f'âœ… Document parsed: {result["word_count"]} words')
            return result
            
        except Exception as e:
            logger.error(f'âŒ Error parsing document: {str(e)}')
            raise
    
    def _validate_file(self, file_path: str) -> bool:
        path = Path(file_path)
        return path.exists() and path.suffix.lower() in self.supported_formats
    
    def add_comments_to_document(self, original_path: str, compliance_results: Dict[str, Any], red_flags: List[Dict[str, Any]], output_path: str) -> str:
        logger.info(f'í³ Adding comments to: {original_path}')
        
        try:
            doc = Document(original_path)
            os.makedirs(os.path.dirname(output_path), exist_ok=True)
            
            # Add header with review summary
            header_para = doc.paragraphs[0]
            header_run = header_para.insert_paragraph_before().add_run('í¿›ï¸ ADGM COMPLIANCE REVIEW')
            header_run.bold = True
            header_run.font.size = Pt(16)
            
            # Add summary
            summary_para = doc.paragraphs[1]
            summary_text = f'Review Date: {datetime.now().strftime("%Y-%m-%d")}\nIssues Found: {len(red_flags)}\n'
            summary_para.insert_paragraph_before().add_run(summary_text)
            
            doc.save(output_path)
            logger.info(f'âœ… Marked-up document saved: {output_path}')
            return output_path
            
        except Exception as e:
            logger.error(f'âŒ Error adding comments: {str(e)}')
            return original_path
